"""端到端集成测试：md → docx 的样式实际生效。

需要真实 pandoc 和 python-docx。默认通过 pytest marker 'integration' 跳过，
显式运行：pytest tests/test_integration_styled.py -m integration -v
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest

pytestmark = pytest.mark.integration

try:
    from docx import Document
except ImportError:
    pytest.skip("python-docx 未安装", allow_module_level=True)

from docx.oxml.ns import qn

from md2doc import converter, pandoc

FIXTURE = Path(__file__).resolve().parent / "fixtures" / "simple.md"


@pytest.fixture(scope="module")
def generated_docx(tmp_path_factory):
    if pandoc.get_version() is None:
        pytest.skip("pandoc 未安装")
    out = tmp_path_factory.mktemp("out") / "result.docx"
    converter.convert_file(FIXTURE, out, "docx", no_mermaid=True, styled=True)
    return Document(str(out))


def _east_asia(style) -> str | None:
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        return None
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        return None
    return rFonts.get(qn("w:eastAsia"))


def test_headings_use_heiti(generated_docx):
    """至少一个 Heading 段落字体为黑体。"""
    heading_styles = {
        s.name for s in generated_docx.styles
        if s.name and s.name.startswith("Heading")
    }
    found_heiti = False
    for name in heading_styles:
        s = generated_docx.styles[name]
        if _east_asia(s) == "黑体":
            found_heiti = True
            break
    assert found_heiti, "未找到 eastAsia=黑体 的 Heading 样式"


def test_image_caption_and_table_caption_styles_exist(generated_docx):
    """生成的 docx 应包含 Image Caption 和 Table Caption 样式。"""
    names = {s.name for s in generated_docx.styles if s.name}
    assert "Image Caption" in names
    assert "Table Caption" in names


def test_number_sections_applied(generated_docx):
    """pandoc --number-sections 应在标题前注入数字。simple.md 中应能找到形如 '1 xxx' 的标题段落。"""
    found = False
    for p in generated_docx.paragraphs:
        if (
            p.style
            and p.style.name
            and p.style.name.startswith("Heading")
            and re.match(r"^\d+(\.\d+)*\s", p.text)
        ):
            found = True
            break
    assert found, f"未找到自动编号的标题，所有标题文本：{[p.text for p in generated_docx.paragraphs if p.style and p.style.name and p.style.name.startswith('Heading')]}"
